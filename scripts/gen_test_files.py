"""生成两个测试文件：制度清单Excel + 制度原文docx + 流程清单Excel"""
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUT = Path(__file__).parent.parent / "data" / "test_files"
OUT.mkdir(parents=True, exist_ok=True)

# ── 数据定义 ──────────────────────────────────────────────

# 模拟"真实企业"上传的非标准格式：
#   - 列名用中文，且和系统字段不一致（"制度编号" ≠ "id"）
#   - 类别名是企业自己的叫法（如"治理类"），不是系统标准的"法人治理"
#   - 流程字段是箭头分隔的纯文本，不是结构化数组
# 系统应该靠 AI 做列名映射和内容理解，把这些数据"翻译"成标准格式。

REGS = [
    {"制度编号": "ZD-001", "制度名称": "公司章程",             "所属类别": "治理类",   "二级分类": "法人治理",   "主责部门": "董事会办公室", "文号": "XX-2024-001", "状态": "现行有效"},
    {"制度编号": "ZD-002", "制度名称": "三重一大决策管理办法",   "所属类别": "治理类",   "二级分类": "决策机制",   "主责部门": "综合办公室",   "文号": "XX-2024-002", "状态": "现行有效"},
    {"制度编号": "ZD-003", "制度名称": "员工招聘与录用管理规定", "所属类别": "人力资源", "二级分类": "招聘管理",   "主责部门": "人力资源部",   "文号": "XX-2024-003", "状态": "现行有效"},
    {"制度编号": "ZD-004", "制度名称": "薪酬福利管理办法",       "所属类别": "人力资源", "二级分类": "薪酬体系",   "主责部门": "人力资源部",   "文号": "XX-2024-004", "状态": "现行有效"},
    {"制度编号": "ZD-005", "制度名称": "财务管理基本制度",       "所属类别": "财务",     "二级分类": "财务核算",   "主责部门": "财务部",       "文号": "XX-2024-005", "状态": "现行有效"},
    {"制度编号": "ZD-006", "制度名称": "固定资产管理办法",       "所属类别": "财务",     "二级分类": "资产管理",   "主责部门": "财务部",       "文号": "XX-2024-006", "状态": "现行有效"},
    {"制度编号": "ZD-007", "制度名称": "合同管理办法",           "所属类别": "法务合规", "二级分类": "合同审查",   "主责部门": "法务部",       "文号": "XX-2024-007", "状态": "现行有效"},
    {"制度编号": "ZD-008", "制度名称": "采购管理办法",           "所属类别": "供应链",   "二级分类": "采购流程",   "主责部门": "采购部",       "文号": "XX-2024-008", "状态": "现行有效"},
    {"制度编号": "ZD-009", "制度名称": "信息安全管理制度",       "所属类别": "信息化",   "二级分类": "信息安全",   "主责部门": "信息技术部",   "文号": "XX-2024-009", "状态": "试运行"},
    {"制度编号": "ZD-010", "制度名称": "内部审计管理办法",       "所属类别": "监督",     "二级分类": "审计制度",   "主责部门": "审计部",       "文号": "XX-2024-010", "状态": "现行有效"},
]

FLOWS = [
    {"流程编号": "LC-001", "流程名称": "员工入职审批",         "业务类别": "人力资源", "关联制度": "员工招聘与录用管理规定", "发起人": "人力资源部", "最终审批人": "总经理",     "审批流程": "人力资源部→部门负责人→分管领导→总经理"},
    {"流程编号": "LC-002", "流程名称": "员工离职审批",         "业务类别": "人力资源", "关联制度": "员工招聘与录用管理规定", "发起人": "员工本人",   "最终审批人": "人力资源部", "审批流程": "员工本人→直属领导→人力资源部→财务部"},
    {"流程编号": "LC-003", "流程名称": "费用报销审批",         "业务类别": "财务",     "关联制度": "财务管理基本制度",       "发起人": "报销人",     "最终审批人": "财务部",     "审批流程": "报销人→部门负责人→财务部审核→分管领导→总经理"},
    {"流程编号": "LC-004", "流程名称": "固定资产购置审批",     "业务类别": "财务",     "关联制度": "固定资产管理办法",       "发起人": "申请部门",   "最终审批人": "总经理",     "审批流程": "申请部门→财务部→分管领导→总经理→董事会"},
    {"流程编号": "LC-005", "流程名称": "合同签署审批",         "业务类别": "法务合规", "关联制度": "合同管理办法",           "发起人": "业务部门",   "最终审批人": "总经理",     "审批流程": "业务部门→法务部审查→财务部→分管领导→总经理"},
    {"流程编号": "LC-006", "流程名称": "采购申请审批",         "业务类别": "供应链",   "关联制度": "采购管理办法",           "发起人": "申请部门",   "最终审批人": "总经理",     "审批流程": "申请部门→采购部→财务部→分管领导→总经理"},
    {"流程编号": "LC-007", "流程名称": "用章申请审批",         "业务类别": "行政",     "关联制度": "公司章程",               "发起人": "申请人",     "最终审批人": "综合办公室", "审批流程": "申请人→部门负责人→综合办公室→分管领导"},
    {"流程编号": "LC-008", "流程名称": "出差审批",             "业务类别": "行政",     "关联制度": "三重一大决策管理办法",   "发起人": "员工本人",   "最终审批人": "部门负责人", "审批流程": "员工本人→直属领导→部门负责人→分管领导"},
    {"流程编号": "LC-009", "流程名称": "IT系统访问权限申请",   "业务类别": "信息化",   "关联制度": "信息安全管理制度",       "发起人": "申请人",     "最终审批人": "信息技术部", "审批流程": "申请人→部门负责人→信息安全专员→信息技术部"},
    {"流程编号": "LC-010", "流程名称": "内部审计整改报告审批", "业务类别": "监督",     "关联制度": "内部审计管理办法",       "发起人": "被审计部门", "最终审批人": "总经理",     "审批流程": "被审计部门→审计部确认→分管领导→总经理"},
]

# ── 1. 制度清单 Excel ───────────────────────────────────

def make_reg_excel():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "制度清单"

    # 样式
    hdr_fill = PatternFill("solid", fgColor="1E3A5F")
    hdr_font = Font(bold=True, color="FFFFFF", size=11)
    alt_fill = PatternFill("solid", fgColor="EFF3F9")
    thin = Border(
        left=Side(style="thin", color="D0D7E3"),
        right=Side(style="thin", color="D0D7E3"),
        top=Side(style="thin", color="D0D7E3"),
        bottom=Side(style="thin", color="D0D7E3"),
    )
    center = Alignment(horizontal="center", vertical="center")
    wrap = Alignment(wrap_text=True, vertical="center")

    # 模拟"真实企业"的非标准列名（中文，且和系统字段名不一致）
    headers = ["制度编号", "制度名称", "所属类别", "二级分类", "主责部门", "文号", "状态"]
    col_w   = [12,         28,         14,         16,         16,         16,     12]

    # 标题行
    ws.append(headers)
    for c, (h, w) in enumerate(zip(headers, col_w), 1):
        cell = ws.cell(1, c)
        cell.font = hdr_font
        cell.fill = hdr_fill
        cell.alignment = center
        cell.border = thin
        ws.column_dimensions[cell.column_letter].width = w

    # 数据行
    for i, reg in enumerate(REGS, 2):
        row = [reg[h] for h in headers]
        ws.append(row)
        fill = alt_fill if i % 2 == 0 else None
        for c in range(1, len(headers) + 1):
            cell = ws.cell(i, c)
            if fill:
                cell.fill = fill
            cell.border = thin
            cell.alignment = center if c in (1, 6, 7) else wrap

    ws.row_dimensions[1].height = 28
    for i in range(2, len(REGS) + 2):
        ws.row_dimensions[i].height = 22

    path = OUT / "XX企业制度清单.xlsx"
    wb.save(path)
    print(f"✓ 制度清单 → {path}")

# ── 2. 流程清单 Excel ───────────────────────────────────

def make_auth_excel():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "业务流程表"

    hdr_fill = PatternFill("solid", fgColor="1E3A5F")
    hdr_font = Font(bold=True, color="FFFFFF", size=11)
    alt_fill = PatternFill("solid", fgColor="EFF3F9")
    thin = Border(
        left=Side(style="thin", color="D0D7E3"),
        right=Side(style="thin", color="D0D7E3"),
        top=Side(style="thin", color="D0D7E3"),
        bottom=Side(style="thin", color="D0D7E3"),
    )
    center = Alignment(horizontal="center", vertical="center")
    wrap = Alignment(wrap_text=True, vertical="center")

    # 模拟"真实企业"的非标准列名（流程字段是"A→B→C"纯文本）
    headers = ["流程编号", "流程名称", "业务类别", "关联制度", "发起人", "最终审批人", "审批流程"]
    col_w   = [10,         22,         14,         24,         12,       14,             50]

    ws.append(headers)
    for c, (h, w) in enumerate(zip(headers, col_w), 1):
        cell = ws.cell(1, c)
        cell.font = hdr_font
        cell.fill = hdr_fill
        cell.alignment = center
        cell.border = thin
        ws.column_dimensions[cell.column_letter].width = w

    for i, fl in enumerate(FLOWS, 2):
        row = [fl[h] for h in headers]
        ws.append(row)
        fill = alt_fill if i % 2 == 0 else None
        for c in range(1, len(headers) + 1):
            cell = ws.cell(i, c)
            if fill:
                cell.fill = fill
            cell.border = thin
            cell.alignment = center if c in (1, 3, 5, 6) else wrap

    ws.row_dimensions[1].height = 28
    for i in range(2, len(FLOWS) + 2):
        ws.row_dimensions[i].height = 32

    path = OUT / "XX公司业务流程表.xlsx"
    wb.save(path)
    print(f"✓ 流程清单 → {path}")

# ── 3. 制度原文 docx ────────────────────────────────────

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    return p

def add_article(doc, no, text):
    p = doc.add_paragraph()
    run_no = p.add_run(f"第{no}条  ")
    run_no.bold = True
    p.add_run(text)
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(4)

def make_reg_docx():
    """生成《采购管理办法》作为制度原文示例（对应 id=4-2-1）"""
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # 正文字体
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)

    # ── 封面 ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("XX 企 业")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    add_heading(doc, "采购管理办法", level=1)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("文件编号：XX-2024-008　　　　版本：V1.0\n发布日期：2024年3月1日　　　　实施日期：2024年4月1日")
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()
    doc.add_page_break()

    # ── 第一章 ──
    add_heading(doc, "第一章  总  则", level=2)
    add_article(doc, "一", "为规范公司采购行为，提高采购效率，降低采购成本，防范采购风险，根据国家有关法律法规及公司实际情况，制定本办法。")
    add_article(doc, "二", "本办法适用于公司所有部门及控股子公司的采购活动，包括货物采购、服务采购及工程承包采购。")
    add_article(doc, "三", "采购工作遵循公开透明、公平竞争、诚实信用、物有所值的原则。")

    # ── 第二章 ──
    add_heading(doc, "第二章  采购组织与职责", level=2)
    add_article(doc, "四", "采购部门是公司采购归口管理部门，负责制度制定、供应商管理、合同归档及采购数据分析等工作。")
    add_article(doc, "五", "各业务部门作为采购需求部门，负责提出采购需求、明确技术参数、参与技术评审及验收工作。")
    add_article(doc, "六", "财务部门负责采购预算管理、资金审核及付款审批工作。")
    add_article(doc, "七", "法务部门负责采购合同的法律审查工作，重大合同须经法律顾问复核。")

    # ── 第三章 ──
    add_heading(doc, "第三章  采购方式", level=2)
    add_article(doc, "八", "公司采购分为公开招标、邀请招标、竞争性谈判、询价采购和单一来源采购五种方式。")
    add_article(doc, "九", "采购金额50万元（含）以上的货物或服务，原则上采用公开招标方式。")
    add_article(doc, "十", "采购金额20万元（含）至50万元以下的，采用邀请招标或竞争性谈判方式。")
    add_article(doc, "十一", "采购金额5万元（含）至20万元以下的，采用询价采购方式，须取得不少于3家供应商报价。")
    add_article(doc, "十二", "采购金额5万元以下的零星采购，可采用单一来源方式，由部门负责人审批。")

    # ── 第四章 ──
    add_heading(doc, "第四章  采购审批权限", level=2)
    add_article(doc, "十三", "各层级采购审批权限如下：")
    tbl = doc.add_table(rows=5, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for cell, text in zip(hdr, ["采购金额", "审批层级", "审批人"]):
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True
    data = [
        ("5万元以下", "部门内部", "部门负责人"),
        ("5万元至20万元", "分管领导审批", "分管副总经理"),
        ("20万元至50万元", "总经理审批", "总经理"),
        ("50万元以上", "董事会审批", "董事会"),
    ]
    for row, (a, b, c) in zip(tbl.rows[1:], data):
        row.cells[0].text = a
        row.cells[1].text = b
        row.cells[2].text = c
    doc.add_paragraph()

    # ── 第五章 ──
    add_heading(doc, "第五章  供应商管理", level=2)
    add_article(doc, "十四", "公司建立合格供应商名录，采购部门定期对供应商进行资质审查和绩效评价。")
    add_article(doc, "十五", "新供应商须经采购部门完成资质审查、现场考察后方可纳入合格供应商名录。")
    add_article(doc, "十六", "供应商年度评价结果分为优秀、合格、待改进、淘汰四个等级，连续两年评价为待改进或一次评价为淘汰的，取消合格供应商资格。")

    # ── 第六章 ──
    add_heading(doc, "第六章  附  则", level=2)
    add_article(doc, "十七", "本办法由采购部门负责解释，自发布之日起施行。")
    add_article(doc, "十八", "本办法未尽事宜，按照国家有关法律法规及公司其他相关制度执行。")

    path = OUT / "采购管理办法（测试）.docx"
    doc.save(path)
    print(f"✓ 制度原文 → {path}")


if __name__ == "__main__":
    make_reg_excel()
    make_auth_excel()
    make_reg_docx()
    print("\n全部文件已生成至：", OUT)
