#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成《制度管理平台用户手册》DOCX"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT = "D:/西北投资/制度助理/制度助理-v1.1.0/西北投资制度管理平台用户手册_v1.1.0.docx"

NAVY   = RGBColor(0x1E, 0x3A, 0x5F)
BLUE   = RGBColor(0x25, 0x63, 0xEB)
MUTED  = RGBColor(0x64, 0x74, 0x8B)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xF0, 0xF4, 0xFF)
BLACK  = RGBColor(0x00, 0x00, 0x00)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, color, sz in kwargs.get('borders', []):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(sz))
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(8 if level == 1 else 6)
    run = p.add_run(text)
    if level == 1:
        run.font.size  = Pt(16)
        run.font.bold  = True
        run.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(22)
    elif level == 2:
        run.font.size  = Pt(13)
        run.font.bold  = True
        run.font.color.rgb = BLUE
    else:
        run.font.size  = Pt(12)
        run.font.bold  = True
        run.font.color.rgb = MUTED
    return p


def add_para(doc, text, indent=0, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.space_before = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK
    return p


def add_step(doc, num, title, desc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f"Step {num}  {title}")
    r1.font.size  = Pt(11)
    r1.font.bold  = True
    r1.font.color.rgb = BLUE
    if desc:
        r2 = p.add_run(f"\n        {desc}")
        r2.font.size  = Pt(11)
        r2.font.color.rgb = MUTED


def add_tip(doc, text, label="提示"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.6)
    r1 = p.add_run(f"▶ {label}：")
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)


def add_warn(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.6)
    r1 = p.add_run("⚠ 注意：")
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)


def add_bullet(doc, items, indent=1):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent  = Cm(indent)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(item)
        run.font.size  = Pt(11)
        run.font.color.rgb = BLACK


def feature_table(doc, rows):
    table = doc.add_table(rows=len(rows)+1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ['功能模块', '主要功能', '适用场景']
    widths = [Cm(3.2), Cm(7.5), Cm(5.5)]
    for i, (h, w) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[i]
        cell.width = w
        set_cell_bg(cell, '1E3A5F')
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold  = True
        run.font.color.rgb = WHITE
        run.font.size  = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, (col1, col2, col3) in enumerate(rows):
        for ci, text in enumerate([col1, col2, col3]):
            cell = table.rows[ri+1].cells[ci]
            if ri % 2 == 0:
                set_cell_bg(cell, 'F0F4FF')
            p = cell.paragraphs[0]
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.space_before = Pt(3)
            run = p.add_run(text)
            run.font.size = Pt(10.5)
    return table


doc = Document()

# ── 页面边距 ──
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── 封面 ──
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("西北投资（园区运营公司）")
r.font.size = Pt(15); r.font.color.rgb = MUTED; r.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("制度管理平台")
r.font.size = Pt(30); r.font.color.rgb = NAVY; r.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("用户使用手册")
r.font.size = Pt(20); r.font.color.rgb = BLUE; r.font.bold = True

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("版本  v1.1.0    |    2026 年 5 月")
r.font.size = Pt(12); r.font.color.rgb = MUTED

doc.add_page_break()

# ─────────────────────────────────────────
# 第一章 产品简介
# ─────────────────────────────────────────
add_heading(doc, "一、产品简介")

add_para(doc, '西北投资制度管理平台是专为西北投资（园区运营公司）内部使用而设计的规章制度与权责清单一体化查阅工具。平台将公司现行全部制度文件、权责审批事项统一数字化收录，支持按职能域、制度大类、关键词等多种方式快速检索，实现“随时查、随时用”。')

doc.add_paragraph()
add_para(doc, "【核心特点】")
add_bullet(doc, [
    "无需安装、无需联网：解压后直接双击 index.html 即可使用，Windows 与 Mac 均支持",
    "覆盖全面：收录公司制度库全部规章制度及权责审批事项",
    "检索高效：支持关键词全文搜索，支持职能域、审批层级、发起部门等多维筛选",
    "阅读友好：制度全文分章节目录导航，支持一键跳转，并自动标注当前阅读位置",
])

# ─────────────────────────────────────────
# 第二章 快速开始
# ─────────────────────────────────────────
add_heading(doc, "二、快速开始（3 步上手）")

add_heading(doc, "2.1  获取文件", level=2)
add_para(doc, "向数字化专员索取 制度助理-v1.1.0.zip 文件，或通过企业内部共享渠道下载。文件大小约 2.1 MB，下载极快。")

add_heading(doc, "2.2  解压文件", level=2)
add_step(doc, 1, "（Windows）", "右键点击 ZIP 文件 → 选择「全部解压缩」→ 点击「提取」，会在同一目录生成「制度助理-v1.1.0」文件夹。")
add_step(doc, 1, "（Mac）", "双击 ZIP 文件，系统自动解压生成「制度助理-v1.1.0」文件夹。")
add_tip(doc, "解压后不要移动文件夹内的任何文件，必须保持 assets 文件夹与 index.html 在同一目录。")

add_heading(doc, "2.3  打开使用", level=2)
add_step(doc, 2, "进入文件夹", "打开「制度助理-v1.1.0」文件夹。")
add_step(doc, 3, "双击 index.html", "系统会自动用默认浏览器打开平台首页。")
add_tip(doc, "推荐使用 Chrome 或 Edge 浏览器，显示效果最佳。Firefox 亦可正常使用。")
add_warn(doc, "Safari 浏览器因安全策略可能有部分功能受限，Mac 用户建议改用 Chrome。")

# ─────────────────────────────────────────
# 第三章 功能模块详解
# ─────────────────────────────────────────
add_heading(doc, "三、功能模块详解")

add_heading(doc, "3.1  系统概览（首页）", level=2)
add_para(doc, "打开平台后进入系统概览页面，提供全局数据视图与快速入口。")
doc.add_paragraph()

feature_table(doc, [
    ("统计卡片", "制度总数、正式发布数量、试运行数量、权责事项总数，一屏掌握全局",  "快速了解平台收录规模"),
    ("制度库看板", "按职能域分组展示制度数量，点击可直接跳转对应域",                "从宏观维度找制度"),
    ("权责看板",   "按职能域分组展示权责事项数量，点击可进入权责清单对应筛选",      "快速定位相关审批流程"),
    ("组织架构图", "西北投资公司全部职能部门与下属项目公司一览，点击部门可跳转权责搜索", "了解组织结构与权责归属"),
])

doc.add_paragraph()
add_tip(doc, "点击统计卡片或看板中的任意域名，可一键跳转至该域的制度或权责列表。")

add_heading(doc, "3.2  制度库", level=2)
add_para(doc, "制度库按照「职能域 → 大类模块 → 子类 → 具体制度」四级结构组织，支持逐级浏览或直接搜索。")
doc.add_paragraph()
add_heading(doc, "进入方式", level=3)
add_bullet(doc, [
    "点击顶部导航「制度库」，或点击首页「制度库看板」中任意职能域",
    "左侧导航栏也常驻有制度库目录树",
])

add_heading(doc, "浏览层级", level=3)
add_bullet(doc, [
    "第一层：7 个职能域总览，展示各域制度数量及简介",
    "第二层：点击职能域后，展示该域下全部大类模块",
    "第三层：点击大类后，按子类分组列出所有具体制度",
    "点击任意制度条目，进入制度详情页查看全文",
])

add_heading(doc, "面包屑导航", level=3)
add_para(doc, "页面顶部始终显示「首页 › 制度库 › 职能域名 › 大类名」路径，点击任意层级可快速返回。")

add_heading(doc, "3.3  制度详情", level=2)
add_para(doc, "制度详情页提供完整的制度文本阅读体验，包含目录导航和授权映射两个辅助功能。")
doc.add_paragraph()

add_heading(doc, "页面布局", level=3)
add_bullet(doc, [
    "左侧：固定浮动目录（TOC），列出制度的章节结构，点击可跳转对应段落，当前阅读位置自动高亮",
    "右侧：制度正文区域，包含标题、文号、各章节内容、附件目录",
    "顶部：「制度原文」/「授权映射」两个选项卡切换",
])

add_heading(doc, "授权映射功能", level=3)
add_para(doc, "切换至「授权映射」选项卡，可查看该制度所关联的权责审批事项列表。点击任一事项，可快速跳转至权责清单中对应的审批流程卡片，无需手动搜索。")

add_tip(doc, "目录会随阅读滚动自动高亮当前章节，快速定位当前位置。")

add_heading(doc, "3.4  权责清单", level=2)
add_para(doc, "权责清单收录公司全部授权审批事项，详细展示每项事务的审批层级、流转步骤和参与部门。")
doc.add_paragraph()

add_heading(doc, "筛选功能", level=3)
feature_table(doc, [
    ("职能域筛选", "顶部彩色标签，按职能域范围筛选",            "只看某个职能域的权责"),
    ("大类 Tab",   "全部 / 项目公司事项 / 西北投资事项",        "区分不同层级的审批事项"),
    ("关键词搜索", "输入名称、发起部门、参与部门、系统名等关键词", "快速找到某个具体事项"),
    ("审批层级筛选","项目公司审批 / 西北投资审批 / 决策机构 / 总部及以上", "按最终审批层级过滤"),
])
doc.add_paragraph()

add_heading(doc, "卡片内容说明", level=3)
add_bullet(doc, [
    "编号徽标：蓝色背景的事项编号（如 A-01），对应权责手册原始编号",
    "发起部门：该事项由哪个部门发起提报",
    "审批流程：以气泡流程图展示，包含每个环节的角色、审批类型、是否为最终审批节点",
    "备注：如有特殊说明或授权限制，以橙色文字标注",
])

# ─────────────────────────────────────────
# 第四章 搜索功能
# ─────────────────────────────────────────
add_heading(doc, "四、搜索功能使用技巧")

add_heading(doc, "4.1  全局搜索（概览页）", level=2)
add_para(doc, "在首页左侧导航栏顶部搜索框输入关键词，可跨制度库和权责清单进行全局搜索。")
add_bullet(doc, [
    "支持制度名称、制度编号（如 GF-001）搜索",
    "结果按制度大类模块分组显示",
    "点击结果条目直接跳转制度详情",
    "按 Esc 键或清除按钮可清空搜索",
])

add_heading(doc, "4.2  权责清单搜索", level=2)
add_para(doc, "在权责清单页面的搜索框，支持以下搜索范围：")
add_bullet(doc, [
    "事项名称（如：投资项目、资产处置）",
    "发起部门（如：财务资金部、经营管理中心）",
    "涉及部门（如：法律风控部）",
    "所用系统名称（如：OA 系统）",
    "流程中的审批角色",
])

add_tip(doc, "搜索词会在结果中以黄色高亮显示，方便快速定位匹配内容。")

add_heading(doc, "4.3  常用搜索示例", level=2)
feature_table(doc, [
    ("找某制度全文",     "在搜索框输入制度名称关键词，如「薪酬」「采购」",      "制度库全局搜索"),
    ("查某部门审批事项", "在权责清单搜索「财务资金部」或「综合办公室」",         "权责清单关键词搜索"),
    ("查某层级审批",     "点击「审批层级」快筛按钮，选择「西北投资审批」",       "权责清单审批层级筛选"),
    ("找某职能域全部制度","点击首页职能域卡片，如「投资管理」",                  "首页看板直接跳转"),
])

# ─────────────────────────────────────────
# 第五章 常见问题
# ─────────────────────────────────────────
add_heading(doc, "五、常见问题解答")

qa_pairs = [
    ("Q：双击 index.html 没有反应？",
     "A：请确认文件夹中 assets 文件夹与 index.html 在同一目录。如果 index.html 被移出了文件夹，单独放置则无法加载样式和数据。"),
    ("Q：打开后页面显示空白或乱码？",
     "A：请尝试换用 Chrome 或 Edge 浏览器打开。部分系统默认浏览器可能不支持现代网页标准。"),
    ("Q：Mac 上打开后有些功能不正常？",
     "A：Mac 建议使用 Chrome 浏览器，Safari 因本地文件安全策略可能限制部分功能。"),
    ("Q：搜索没有结果？",
     "A：请检查关键词是否正确，尝试缩短关键词（如用「投资」代替「投资管理相关事项」）。权责清单的搜索框与制度库搜索框是分开的，请确认在正确页面搜索。"),
    ("Q：制度目录显示不全？",
     "A：部分较短或格式特殊的制度可能只有正文无分章结构，此时目录为空属正常情况。"),
    ("Q：如何分享给同事？",
     "A：将 制度助理-v1.1.0.zip 文件直接发给同事，对方解压后按本手册操作即可，无需安装任何软件。"),
    ("Q：数据多久更新一次？",
     "A：当前版本（v1.1.0）数据截至 2026 年 1 月。如需更新数据，请联系数字化专员获取新版本。"),
]

for q, a in qa_pairs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(q)
    r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.6)
    p2.paragraph_format.space_after = Pt(6)
    r2 = p2.add_run(a)
    r2.font.size = Pt(11); r2.font.color.rgb = MUTED

# ─────────────────────────────────────────
# 第六章 附录
# ─────────────────────────────────────────
add_heading(doc, "六、附录：职能域说明")

add_para(doc, "平台将全部制度和权责事项按以下 7 个职能域分类管理：")
doc.add_paragraph()

domain_table = doc.add_table(rows=8, cols=3)
domain_table.style = 'Table Grid'
domain_table.alignment = WD_TABLE_ALIGNMENT.LEFT
headers2 = ['职能域', '涵盖范围', '典型制度/事项举例']
for i, h in enumerate(headers2):
    cell = domain_table.rows[0].cells[i]
    set_cell_bg(cell, '1E3A5F')
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.font.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

domain_rows = [
    ("战略投资", "项目投资决策、资本运作、并购重组", "投资管理办法、资产处置审批"),
    ("财务管理", "预算、资金、成本、税务、会计核算", "资金审批、费用报销、银行账户管理"),
    ("人力资源", "招聘、薪酬、绩效、培训、劳动关系", "员工招聘审批、薪酬调整"),
    ("行政综合", "行政事务、文件管理、合同管理、档案", "合同签署审批、公章使用"),
    ("法律风控", "法务管理、合规、风险管理、审计纪检", "法律事务委托、违规处理"),
    ("工程建设", "工程招采、施工管理、竣工验收", "招标审批、工程变更审批"),
    ("运营管理", "产业园区运营、招商、物业管理", "招商合同审批、园区服务管理"),
]
for ri, (col1, col2, col3) in enumerate(domain_rows):
    for ci, text in enumerate([col1, col2, col3]):
        cell = domain_table.rows[ri+1].cells[ci]
        if ri % 2 == 0:
            set_cell_bg(cell, 'F0F4FF')
        p = cell.paragraphs[0]
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.space_before = Pt(3)
        run = p.add_run(text)
        run.font.size = Pt(10.5)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("— 如需技术支持，请联系产业发展中心数字化专员 —")
r.font.size = Pt(11); r.font.color.rgb = MUTED

doc.save(OUT)
print(f"✅ 用户手册已生成：{OUT}")
